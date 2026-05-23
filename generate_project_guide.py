from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_SCRIPTS = Path(
    r"C:\Users\rodrigo.neiland\.codex\plugins\cache\openai-primary-runtime\documents\26.515.10909\skills\documents\scripts"
)
sys.path.append(str(SKILL_SCRIPTS))

from table_geometry import apply_table_geometry, exact_column_widths  # noqa: E402


DOC_PATH = Path(
    r"C:\Users\rodrigo.neiland\OneDrive - ESPM\Documentos\3sem\Github\HackathonPrevisao\guia_projeto_estoque_q4_2024.docx"
)

LETTER_WIDTH = Inches(8.5)
LETTER_HEIGHT = Inches(11)
MARGIN = Inches(1)
HEADER_FOOTER_DISTANCE = Inches(0.492)

TITLE_COLOR = RGBColor(11, 37, 69)
HEADING_BLUE = RGBColor(46, 116, 181)
HEADING_DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(95, 99, 104)
LIGHT_FILL = "E8EEF5"
VERY_LIGHT_FILL = "F4F6F9"
TABLE_GRID = "CAD3DD"


def set_rfonts(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def style_run(run, *, font="Calibri", size=11, color=None, bold=None, italic=None) -> None:
    set_rfonts(run, font)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_section(section) -> None:
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = LETTER_WIDTH
    section.page_height = LETTER_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = HEADER_FOOTER_DISTANCE
    section.footer_distance = HEADER_FOOTER_DISTANCE


def set_paragraph_spacing(paragraph, *, before=0, after=0, line=1.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = TITLE_COLOR
    title_pf = title.paragraph_format
    title_pf.space_before = Pt(0)
    title_pf.space_after = Pt(4)
    title_pf.line_spacing = 1.0

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.italic = False
    subtitle.font.color.rgb = MUTED
    sub_pf = subtitle.paragraph_format
    sub_pf.space_before = Pt(0)
    sub_pf.space_after = Pt(14)
    sub_pf.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, HEADING_BLUE),
        ("Heading 2", 13, 14, 7, HEADING_BLUE),
        ("Heading 3", 12, 10, 5, HEADING_DARK),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.0


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_borders(cell, color: str = TABLE_GRID) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side in ("top", "left", "bottom", "right"):
        el = tc_borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), color)


def set_table_borders(table, color: str = TABLE_GRID) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), color)


def add_bottom_rule(paragraph, color: str = "D7DBE2") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    style_run(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_header_footer(section) -> None:
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header_p, before=0, after=0, line=1.0)
    hr = header_p.add_run("Guia do projeto | Planejamento de estoque")
    style_run(hr, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer_p, before=0, after=0, line=1.0)
    fr = footer_p.add_run("Página ")
    style_run(fr, size=9, color=MUTED)
    add_page_field(footer_p)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    set_paragraph_spacing(p, before=0, after=6, line=1.25)
    run = p.add_run(text)
    style_run(run, size=11)


def add_formula_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    set_paragraph_spacing(p, before=0, after=4, line=1.0)
    run = p.add_run(text)
    style_run(run, font="Consolas", size=10, color=TITLE_COLOR)


def add_table(doc: Document, headers, rows, widths_dxa) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        cell = hdr[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, LIGHT_FILL)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=0, after=0, line=1.0)
        run = p.add_run(text)
        style_run(run, size=10.5, bold=True, color=TITLE_COLOR)

    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            if row_idx % 2 == 1:
                shade_cell(cell, VERY_LIGHT_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=0, after=0, line=1.15)
            run = p.add_run(value)
            style_run(run, size=10.25)

    apply_table_geometry(table, exact_column_widths(widths_dxa))
    set_table_borders(table)
    doc.add_paragraph()


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=4, line=1.0)
    run = p.add_run("GUIA PRÁTICO DO PROJETO")
    style_run(run, size=11, color=MUTED, bold=True)

    title = doc.add_paragraph("Planejamento de estoque para o Q4/2024", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph(
        "O que o desafio pede, como executar e quais soluções fazem mais sentido",
        style="Subtitle",
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph(style="Normal")
    set_paragraph_spacing(meta, before=0, after=10, line=1.15)
    text = meta.add_run(
        "Base do repositório analisada em 22/05/2026 | 29 SKUs | 2 lojas | categoria gripe e resfriado"
    )
    style_run(text, size=10, color=MUTED)

    rule = doc.add_paragraph()
    set_paragraph_spacing(rule, before=0, after=8, line=1.0)
    add_bottom_rule(rule)


def build_document() -> Document:
    doc = Document()
    style_document(doc)
    add_header_footer(doc.sections[0])
    add_title_block(doc)

    doc.add_heading("Resumo executivo", level=1)
    add_body_paragraph(
        doc,
        "Pelo material existente neste repositório, a avaliação formal do desafio não cobre o segundo semestre inteiro. Ela cobre o último trimestre de 2024, de 1 de outubro de 2024 a 31 de dezembro de 2024. O histórico sugerido para treino vai de 2 de janeiro de 2023 a 30 de setembro de 2024.",
    )
    add_body_paragraph(
        doc,
        "Na prática, vocês precisam propor uma política de reposição para cada par SKU-loja, simular dia a dia o que teria acontecido no Q4/2024 e provar que essa política melhora o equilíbrio entre nível de serviço, capital médio em estoque e custo operacional de pedir reposição.",
    )
    add_body_paragraph(
        doc,
        "O caminho mais seguro para o projeto é tratar isso como um problema clássico de inventário com previsão de demanda e política (s, S): prever demanda por SKU-loja, transformar essa previsão em ponto de reposição e estoque-alvo, e validar tudo com um backtest sem vazamento temporal.",
    )

    doc.add_heading("O que precisa ser entregue", level=1)
    doc.add_heading("Entregável 1: policy.csv", level=2)
    add_body_paragraph(
        doc,
        "Criem um arquivo com uma linha por SKU e loja, usando pelo menos as colunas product, location, reorder_point_s e order_up_to_S. Se vocês decidirem usar outra política, como revisão periódica ou base-stock, documentem claramente a regra e mantenham a lógica consistente no simulador.",
    )
    doc.add_heading("Entregável 2: simulador", level=2)
    add_body_paragraph(
        doc,
        "O simulador precisa rodar dia a dia entre 1 de outubro de 2024 e 31 de dezembro de 2024. Ele deve receber a política escolhida, respeitar lead time, dar baixa na demanda real do período e calcular os KPIs finais. Este é o coração técnico do trabalho, porque é ele que prova se a política é melhor ou não.",
    )
    doc.add_heading("Entregável 3: pitch", level=2)
    add_body_paragraph(
        doc,
        "Nos slides, o mais importante é mostrar o raciocínio de negócio: por que a política de vocês reduz ruptura sem estourar capital parado, como o backtest foi montado e qual foi o ganho contra o baseline da casa. Um único SKU bem narrado costuma valer mais do que muitas fórmulas jogadas na tela.",
    )

    doc.add_heading("Como ler a base de dados", level=1)
    add_body_paragraph(
        doc,
        "A leitura abaixo resume quais arquivos realmente entram na solução e como cada um ajuda na modelagem.",
    )
    add_table(
        doc,
        headers=["Arquivo", "Usar para", "Grão", "Observação"],
        rows=[
            ["1_hierarquia.csv", "Validar a categoria", "Categoria", "Confirma que o escopo é gripe e resfriado."],
            ["2_produtos_locais.csv", "Custos e cadastro", "SKU x local", "Traz purchase_price, sales_price, minimum_delivery_batch e cost_of_ordering."],
            ["3_locais.csv", "Lead time da loja", "Local", "Mostra o CD, o tipo de local e o lead time CD-loja."],
            ["4_campanhas.csv", "Calendário promocional", "Campanha", "Permite marcar períodos promocionais."],
            ["5_produtos_locais_campanhas.csv", "Vínculo promoção-SKU", "SKU x loja x campanha", "Diz quais produtos entraram em quais campanhas."],
            ["6_fornecedores.csv", "Restrição avançada", "SKU x fornecedor x local", "Útil se vocês quiserem modelar o elo fornecedor-CD como bônus."],
            ["7_saldo.csv", "Estoque e capital", "SKU x loja x dia", "É a foto diária do saldo e ajuda a iniciar o backtest."],
            ["8_inventario_venda.csv", "Demanda e eventos", "SKU x loja x dia x movimento", "SALE é a base principal de demanda; os demais tipos ajudam a entender a operação real."],
        ],
        widths_dxa=[2088, 2232, 1440, 3600],
    )
    add_body_paragraph(
        doc,
        "O enunciado menciona um dicionário detalhado chamado Estrutura dos Dados.xlsx, mas esse arquivo não apareceu no repositório atual. Se ele existir fora desta pasta, vale consultar para validar convenções de saldo e movimentos.",
    )

    doc.add_heading("Sinais importantes já visíveis na base", level=1)
    doc.add_heading("As duas lojas não podem receber a mesma política", level=2)
    add_body_paragraph(
        doc,
        "No Q4/2024, a loja 1314 concentrou aproximadamente 1.051 unidades vendidas, enquanto a loja 841 ficou perto de 174. Além disso, os lead times são muito diferentes: 3 dias para a loja 841 e 9 dias para a loja 1314. Isso praticamente obriga parâmetros separados por loja.",
    )
    doc.add_heading("Campanha importa", level=2)
    add_body_paragraph(
        doc,
        "Cerca de 30% das unidades vendidas no Q4/2024 aconteceram em dias ligados a campanhas promocionais. Mesmo uma regra simples de uplift promocional já pode melhorar o forecast e evitar subestimação da demanda em datas especiais.",
    )
    doc.add_heading("A demanda é concentrada", level=2)
    add_body_paragraph(
        doc,
        "O SKU 18064 domina o volume da categoria, então uma análise ABC faz sentido. Vocês não precisam usar modelos diferentes por SKU, mas vale dedicar mais atenção aos itens A e não assumir que todos têm o mesmo perfil de giro.",
    )
    doc.add_heading("O lote mínimo da loja não é o maior problema", level=2)
    add_body_paragraph(
        doc,
        "Nas 58 linhas de loja em 2_produtos_locais.csv, o campo minimum_delivery_batch aparece como 1. Isso indica que, no nível loja, o gargalo principal tende a ser lead time, variabilidade de demanda e custo de pedido, não restrição pesada de lote.",
    )

    doc.add_heading("Fluxo recomendado de execução", level=1)
    for title, body in [
        (
            "Passo 1: montar um painel diário SKU-loja",
            "Criem uma grade diária por SKU e loja entre 2 de janeiro de 2023 e 31 de dezembro de 2024. Esse painel vira a tabela mestre para unir saldo, vendas, campanhas, preço de compra e lead time.",
        ),
        (
            "Passo 2: definir a demanda observada",
            "Usem os movimentos com type = SALE como demanda principal. Vendas são negativas no arquivo, então a quantidade vendida precisa ser invertida para ficar positiva. Os demais movimentos podem entrar no simulador realista, mas não são a melhor base para o forecast de consumo do cliente.",
        ),
        (
            "Passo 3: tratar ruptura como demanda censurada",
            "Quando o estoque zera, a venda observada pode ficar menor que a demanda real, porque parte do cliente foi perdida. Uma heurística defensável é substituir a demanda de dias em ruptura pelo máximo entre a venda observada e uma média recente de dias equivalentes sem ruptura.",
        ),
        (
            "Passo 4: construir o forecast por SKU-loja",
            "Comecem com algo simples e explicável: média móvel ponderada, média das últimas semanas equivalentes ou média por dia da semana com janela recente. Só depois avaliem Prophet, ARIMA ou LightGBM, porque o ganho desses modelos depende muito mais da calibração do que do nome do algoritmo.",
        ),
        (
            "Passo 5: converter previsão em política de estoque",
            "Depois do forecast, traduzam a demanda esperada em parâmetros de reposição. O ponto de pedido deve cobrir a demanda durante o lead time mais um estoque de segurança; o nível S deve cobrir o reabastecimento até a próxima revisão sem inflar demais o capital parado.",
        ),
        (
            "Passo 6: montar o simulador com convenção clara",
            "Escolham uma convenção operacional e não mudem no meio do caminho. A mais simples é: receber pedidos previstos para o dia no início da manhã, decidir novo pedido com base na posição de estoque, consumir a demanda do dia e fechar o saldo para o próximo dia.",
        ),
        (
            "Passo 7: iterar pelo score final, não só pelo forecast",
            "O melhor forecast nem sempre gera a melhor política. No hackathon, o que importa é o score do backtest com nível de serviço acima de 92%. Então vale calibrar s, S e o fator de segurança diretamente em cima do simulador.",
        ),
    ]:
        doc.add_heading(title, level=2)
        add_body_paragraph(doc, body)

    doc.add_heading("Como parametrizar uma política (s, S)", level=1)
    add_body_paragraph(
        doc,
        "Uma forma defensável de começar é usar a mesma família de política para todos os itens e mudar apenas os parâmetros por SKU-loja. Considerem L como lead time da loja, mu como demanda média diária prevista e sigma como a variabilidade diária estimada.",
    )
    add_formula_paragraph(doc, "mu_LT = mu * L")
    add_formula_paragraph(doc, "safety_stock = z * sigma * sqrt(L)")
    add_formula_paragraph(doc, "s = ceil(mu_LT + safety_stock)")
    add_formula_paragraph(doc, "S = ceil(s + mu * review_cover_days)")
    add_body_paragraph(
        doc,
        "A loja 841 deve usar L = 3 e a loja 1314 deve usar L = 9. O parâmetro z pode ser calibrado por backtest, e review_cover_days pode representar a cobertura desejada até a próxima revisão do estoque. Se quiserem simplificar, testem uma grade curta de valores para z e para dias de cobertura em vez de tentar otimizar tudo de uma vez.",
    )
    add_body_paragraph(
        doc,
        "No simulador, o pedido do dia deve ser calculado pela posição de estoque, não só pelo saldo físico. Em outras palavras, considerem estoque em mãos mais pedidos em trânsito. Isso evita pedir duas vezes pela mesma necessidade.",
    )

    doc.add_heading("Soluções possíveis", level=1)
    add_table(
        doc,
        headers=["Abordagem", "O que entra", "Vantagem", "Risco principal"],
        rows=[
            [
                "MVP defensável",
                "Média móvel por SKU-loja, desvio-padrão histórico e política (s, S) simples.",
                "Rápida, explicável e boa para sair do zero.",
                "Pode errar promoção, ruptura censurada e caudas de demanda.",
            ],
            [
                "Intermediária recomendada",
                "Forecast por loja, ajuste por campanha, correção simples de censura e grid search de s e S.",
                "Mantém explicabilidade e costuma performar melhor no score final.",
                "Exige mais cuidado com o simulador e a calibração.",
            ],
            [
                "Avançada",
                "Modelo probabilístico ou de machine learning, segmentação ABC e simulação fornecedor-CD-loja.",
                "Captura mais nuances e pode render bônus.",
                "Aumenta muito o risco de bug, overfitting e falta de tempo.",
            ],
        ],
        widths_dxa=[1728, 3312, 2088, 2232],
    )

    doc.add_heading("Minha recomendação direta", level=1)
    add_body_paragraph(
        doc,
        "Se o objetivo é entregar algo forte, defendível e com boa chance de funcionar, eu seguiria a abordagem intermediária recomendada. Ela é sofisticada o suficiente para capturar diferenças entre lojas, campanhas e ruptura censurada, mas ainda simples o bastante para ser explicada em cinco minutos.",
    )
    add_body_paragraph(
        doc,
        "A receita prática seria esta: forecast diário por SKU-loja usando janela recente e componente por dia da semana; uplift promocional quando o SKU estiver em campanha; correção leve de demanda censurada em dias de ruptura; política (s, S) com parâmetros calibrados por backtest; e segmentação ABC para dar mais atenção aos itens de maior giro.",
    )
    add_body_paragraph(
        doc,
        "Se o tempo apertar, não troquem o simulador por um modelo mais chique. Um backtest confiável com forecast simples costuma ganhar de uma previsão sofisticada sem consistência operacional.",
    )

    doc.add_heading("Regras do simulador que não podem falhar", level=1)
    doc.add_heading("Estado inicial", level=2)
    add_body_paragraph(
        doc,
        "Iniciem o Q4 usando o último saldo disponível antes de 1 de outubro de 2024, de preferência o snapshot de 30 de setembro de 2024 em 7_saldo.csv. Isso reduz arbitrariedade e aproxima a simulação da operação real.",
    )
    doc.add_heading("Recebimento e lead time", level=2)
    add_body_paragraph(
        doc,
        "Todo pedido emitido no dia d só pode entrar no estoque em d + lead_time. Não usem informação posterior ao dia da decisão e não permitam reposição instantânea. Esse é o principal ponto de risco metodológico do trabalho.",
    )
    doc.add_heading("Ruptura e venda perdida", level=2)
    add_body_paragraph(
        doc,
        "Quando a demanda do dia for maior que o saldo disponível, o estoque deve parar em zero e a diferença deve virar venda perdida. Esses dias contam contra o nível de serviço, então vale registrar explicitamente quando a ruptura aconteceu.",
    )
    doc.add_heading("KPIs finais", level=2)
    add_body_paragraph(
        doc,
        "O projeto será julgado por nível de serviço, capital médio em estoque e custo total de reposição. O nível de serviço precisa ficar acima de 92%, então não faz sentido perseguir estoque mínimo se isso derrubar a disponibilidade do item para o cliente.",
    )

    doc.add_heading("O que mostrar no pitch", level=1)
    for title, body in [
        (
            "Mensagem principal",
            "Abram com uma frase simples: qual política vocês propõem e por que ela melhora o trade-off entre ruptura e capital parado.",
        ),
        (
            "Pipeline analítico",
            "Mostrem em um slide o fluxo dados -> forecast -> política -> simulador -> KPI. Isso deixa claro que o trabalho não foi só um chute de parâmetros.",
        ),
        (
            "Um SKU bem contado",
            "Escolham um item representativo, mostrem o gráfico do estoque real versus o estoque simulado no Q4/2024 e apontem onde a política de vocês teria evitado excesso ou ruptura.",
        ),
        (
            "Próximos passos honestos",
            "Fechem dizendo o que melhorariam com mais tempo: modelagem da cadeia completa, tratamento mais robusto de censura, melhor ajuste de campanha ou otimização por fornecedor.",
        ),
    ]:
        doc.add_heading(title, level=2)
        add_body_paragraph(doc, body)

    doc.add_heading("Plano de ação imediato", level=1)
    add_body_paragraph(
        doc,
        "Se vocês forem começar agora, a ordem mais eficiente é: montar o painel diário, criar um simulador burro para validar a lógica, gerar uma política baseline simples, medir os KPIs, e só então melhorar forecast e calibração. Isso reduz risco e dá uma trilha clara de evolução para o projeto.",
    )
    add_body_paragraph(
        doc,
        "Em resumo: foquem menos em descobrir o modelo perfeito e mais em construir uma política testável, separada por loja, respeitando lead time e sustentada por um backtest honesto. É isso que transforma a análise em recomendação de negócio.",
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    main()
